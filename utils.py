import PyPDF2
import pdfplumber # Better PDF extraction
import streamlit as st
from typing import List
import re
import requests
from bs4 import BeautifulSoup
import pandas as pd
import docx
import io

def extract_text_from_pdf(pdf_file) -> str:
    """Enhanced PDF text extraction with improved handling"""
    try:
        # Try pdfplumber first (better for complex layouts)
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean up common PDF formatting issues
                        page_text = page_text.replace('\n', ' ').replace('\r', ' ')
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {e}")
                    continue

        # If pdfplumber fails, fallback to PyPDF2
        if not text.strip():
            pdf_file.seek(0) # Reset file pointer
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    continue

        return clean_extracted_text(text)
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_website(url: str) -> str:
    """Extract text content from a website URL"""
    try:
        # Add timeout and headers to avoid blocking
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text from main content areas
        content_selectors = ['main', 'article', '.content', '.post', '.entry-content']
        text = ""
        
        for selector in content_selectors:
            elements = soup.select(selector)
            if elements:
                for elem in elements:
                    text += elem.get_text(separator=' ', strip=True) + "\n\n"
                break
        
        # If no main content found, extract from body
        if not text.strip():
            body = soup.find('body')
            if body:
                text = body.get_text(separator=' ', strip=True)
        
        return clean_extracted_text(text)
        
    except Exception as e:
        st.error(f"Error extracting from website: {str(e)}")
        return ""

def extract_text_from_txt(txt_file) -> str:
    """Extract text from TXT file"""
    try:
        # Handle both uploaded file and file path
        if hasattr(txt_file, 'read'):
            content = txt_file.read()
        else:
            with open(txt_file, 'rb') as f:
                content = f.read()
        
        # Try to decode with common encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = content.decode(encoding)
                return clean_extracted_text(text)
            except UnicodeDecodeError:
                continue
        
        return ""
    except Exception as e:
        st.error(f"Error reading TXT file: {str(e)}")
        return ""

def extract_text_from_csv(csv_file) -> str:
    """Extract text from CSV file by converting to readable format"""
    try:
        # Read CSV
        df = pd.read_csv(csv_file)
        
        # Convert to text format
        text = f"CSV Data Summary:\nColumns: {', '.join(df.columns)}\nRows: {len(df)}\n\n"
        
        # Add column descriptions
        for col in df.columns:
            text += f"Column '{col}':\n"
            # Add sample values (first 5 unique)
            sample_values = df[col].dropna().unique()[:5]
            text += f"Sample values: {', '.join(map(str, sample_values))}\n"
            
            # Add basic statistics for numeric columns
            if df[col].dtype in ['int64', 'float64']:
                text += f"Statistics: Min={df[col].min()}, Max={df[col].max()}, Mean={df[col].mean():.2f}\n"
            text += "\n"
        
        # Add first few rows as examples
        text += "Sample Data Rows:\n"
        for idx, row in df.head(10).iterrows():
            row_text = ", ".join([f"{col}: {val}" for col, val in row.items()])
            text += f"Row {idx + 1}: {row_text}\n"
        
        return clean_extracted_text(text)
        
    except Exception as e:
        st.error(f"Error reading CSV file: {str(e)}")
        return ""

def extract_text_from_docx(docx_file) -> str:
    """Extract text from Word document"""
    try:
        doc = docx.Document(docx_file)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            text += "\n--- Table Content ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"
        
        return clean_extracted_text(text)
        
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return ""

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""

    # Fix common extraction issues
    text = re.sub(r'\s+', ' ', text) # Normalize whitespace
    text = re.sub(r'--- Page \d+ ---', '\n\n', text) # Remove page markers

    # Fix decimal number formatting (common PDF issue)
    text = re.sub(r'\$(\d+)\s+(\d+)', r'$\1.\2', text) # Fix "$0 10" -> "$0.10"
    text = re.sub(r'(\d+)\s+-\s*(\d+)', r'\1-\2', text) # Fix "10 - 20" -> "10-20"

    # Remove problematic characters but keep important punctuation
    text = re.sub(r'[^\w\s.,!?;:()\-\'"$%@#/]', '', text)

    # Clean up paragraph structure
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', text) # Fix sentence spacing

    return text.strip()

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    """Enhanced text chunking with better content preservation"""
    if not text or len(text.strip()) < 50:
        return []

    # Split by paragraphs first, then sentences if needed
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""

    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue

        # If current paragraph fits in current chunk
        if len(current_chunk) + len(paragraph) <= chunk_size:
            current_chunk += ("\n\n" if current_chunk else "") + paragraph
        else:
            # Save current chunk if it has content
            if current_chunk.strip():
                chunks.append(current_chunk.strip())

            # Start new chunk
            if len(paragraph) <= chunk_size:
                current_chunk = paragraph
            else:
                # Split long paragraphs by sentences
                sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                temp_chunk = ""
                for sentence in sentences:
                    if len(temp_chunk) + len(sentence) <= chunk_size:
                        temp_chunk += (" " if temp_chunk else "") + sentence
                    else:
                        if temp_chunk:
                            chunks.append(temp_chunk.strip())
                        temp_chunk = sentence
                current_chunk = temp_chunk

    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # Filter out very short chunks and validate content
    valid_chunks = []
    for chunk in chunks:
        # Skip chunks that are mostly metadata or too short
        if (len(chunk.split()) > 15 and
            not chunk.startswith("Chunk") and
            "energy_storage_research" not in chunk.lower()):
            valid_chunks.append(chunk)

    return valid_chunks

def validate_chunk_content(chunk: str) -> bool:
    """Validate that chunk contains meaningful content"""
    # Check for metadata contamination
    problematic_patterns = [
        r'^Chunk \d+',
        r'^[A-Za-z_]+\.pdf',
        r'^\[Source:.*\]$',
        r'^Page \d+$'
    ]

    for pattern in problematic_patterns:
        if re.match(pattern, chunk.strip()):
            return False

    # Ensure minimum content quality
    words = chunk.split()
    return len(words) >= 15 and len([w for w in words if len(w) > 2]) >= 10

def process_multiple_files(uploaded_files) -> List[str]:
    """Process multiple files of different formats"""
    all_chunks = []
    
    for uploaded_file in uploaded_files:
        st.write(f"Processing: {uploaded_file.name}")
        
        # Get file extension
        file_extension = uploaded_file.name.lower().split('.')[-1]
        text = ""
        
        # Extract text based on file type
        if file_extension == 'pdf':
            text = extract_text_from_pdf(uploaded_file)
        elif file_extension == 'txt':
            text = extract_text_from_txt(uploaded_file)
        elif file_extension == 'csv':
            text = extract_text_from_csv(uploaded_file)
        elif file_extension in ['docx', 'doc']:
            if file_extension == 'docx':
                text = extract_text_from_docx(uploaded_file)
            else:
                st.warning(f"DOC format not supported, please convert {uploaded_file.name} to DOCX")
                continue
        else:
            st.warning(f"Unsupported file format: {file_extension}")
            continue
            
        # Process extracted text
        if text and len(text.strip()) > 100:
            chunks = chunk_text(text)
            if chunks:
                labeled_chunks = []
                for i, chunk in enumerate(chunks):
                    if validate_chunk_content(chunk):
                        labeled_chunk = f"[Source: {uploaded_file.name}, Chunk {i+1}]\n{chunk}"
                        labeled_chunks.append(labeled_chunk)
                
                all_chunks.extend(labeled_chunks)
                st.success(f"{uploaded_file.name}: {len(labeled_chunks)} valid chunks created")
            else:
                st.warning(f"{uploaded_file.name}: No valid chunks could be created")
        else:
            st.error(f"Could not extract sufficient text from {uploaded_file.name}")
    
    return all_chunks

def process_website_urls(urls: List[str]) -> List[str]:
    """Process multiple website URLs"""
    all_chunks = []
    
    for url in urls:
        if not url.strip():
            continue
            
        st.write(f"Processing URL: {url}")
        
        text = extract_text_from_website(url.strip())
        
        if text and len(text.strip()) > 100:
            chunks = chunk_text(text)
            if chunks:
                labeled_chunks = []
                for i, chunk in enumerate(chunks):
                    if validate_chunk_content(chunk):
                        labeled_chunk = f"[Source: {url}, Chunk {i+1}]\n{chunk}"
                        labeled_chunks.append(labeled_chunk)
                
                all_chunks.extend(labeled_chunks)
                st.success(f"URL processed: {len(labeled_chunks)} valid chunks created")
            else:
                st.warning(f"No valid chunks could be created from {url}")
        else:
            st.error(f"Could not extract sufficient text from {url}")
    
    return all_chunks

def debug_document_content(documents: List[str], max_chunks: int = 5):
    """Debug function to display document content"""
    st.write("**Document Content Preview:**")
    for i, doc in enumerate(documents[:max_chunks]):
        st.write(f"**Chunk {i}:**")
        # Show first 300 characters
        preview = doc[:300] + "..." if len(doc) > 300 else doc
        st.text(preview)
        st.write("---")
